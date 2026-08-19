#!/usr/bin/env python3
"""
Poisson DiD regressions and Word tables.

Reads docs/panel_data.csv (built by did_analysis.py) and produces:

  Table 1 — DiD summary: mean deaths by danger index tercile,
             pre vs post SFA, with t-tests and High-Low DiD row
  Table 2 — Poisson regression: 3 specifications with Conley (1999)
             spatial HAC standard errors (50 km cutoff)

Output: docs/paper_tables_regressions_<today>.docx

Run did_analysis.py first whenever the danger index or data changes.
"""

import os, sys, warnings
from pathlib import Path

_VENV = Path(__file__).resolve().parent.parent / ".venv"
_PY   = _VENV / "bin" / "python"
if _PY.exists() and Path(sys.prefix).resolve() != _VENV.resolve():
    os.execv(str(_PY), [str(_PY), str(Path(__file__).resolve()), *sys.argv[1:]])

warnings.filterwarnings("ignore")

import numpy as np
import pandas as pd
from datetime import date
from scipy.spatial import cKDTree
from scipy import stats
from pyproj import Transformer
import geopandas as gpd
import statsmodels.api as sm
from docx import Document

import basemap_common as bc
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
REPO     = Path(__file__).resolve().parent.parent
DOCS     = REPO / "docs"
PANEL    = DOCS / "panel_data.csv"
TODAY    = date.today().isoformat()
OUT_DOCX = DOCS / f"paper_tables_regressions_{TODAY}.docx"

# ── Load panel ─────────────────────────────────────────────────────────────────
print("=" * 70)
print("Loading panel...")
print("=" * 70)
if not PANEL.exists():
    raise FileNotFoundError(f"{PANEL} not found — run did_analysis.py first.")

panel = pd.read_csv(PANEL)
print(f"  {len(panel)} rows | {panel['cell_id'].nunique()} cells x 2 periods")

# ── Compute missing columns if panel was built by an older script ──────────────
if "dist_fence_km" not in panel.columns:
    print("  dist_fence_km missing — computing from fence GDB...")
    ped    = gpd.read_file(str(bc.FENCE_GDB), layer=0)
    veh    = gpd.read_file(str(bc.FENCE_GDB), layer=1)
    ped_az = ped[ped["STATE_ABBR"] == "AZ"].to_crs("EPSG:32612")
    veh_az = veh[veh["STATE_ABBR"] == "AZ"].to_crs("EPSG:32612")

    def _extract(geom):
        if geom.geom_type in ("LineString", "LinearRing"): return list(geom.coords)
        pts = []
        for p in geom.geoms: pts.extend(_extract(p))
        return pts

    fence_pts = []
    for geom in pd.concat([ped_az.geometry, veh_az.geometry]):
        fence_pts.extend(_extract(geom))

    fence_tree = cKDTree(np.array(fence_pts))
    cu = panel[["cell_id","lon","lat"]].drop_duplicates().copy()
    gdf = gpd.GeoDataFrame(cu, geometry=gpd.points_from_xy(cu["lon"], cu["lat"]),
                           crs="EPSG:4326").to_crs("EPSG:32612")
    d, _ = fence_tree.query(np.column_stack([gdf.geometry.x, gdf.geometry.y]))
    cu["dist_fence_km"] = d / 1000.0
    panel = panel.merge(cu[["cell_id","dist_fence_km"]], on="cell_id")

for col, a, b in [("Post_x_Di",    "Post", "D_i"),
                  ("Post_x_fence", "Post", "dist_fence_km"),
                  ("Post_x_gap",   "Post", "dist_gap_km")]:
    if col not in panel.columns:
        panel[col] = panel[a] * panel[b]

if "log_apprehensions" not in panel.columns:
    SECTOR_APPREHENSIONS = {
        ("tucson", 0): 431_012, ("tucson", 1): 125_527,
        ("yuma",   0):  84_933, ("yuma",   1):  14_620,
    }
    panel["sector"] = np.where(panel["lon"] < -113.0, "yuma", "tucson")
    panel["log_apprehensions"] = panel.apply(
        lambda r: np.log(SECTOR_APPREHENSIONS[(r["sector"], int(r["Post"]))]), axis=1)

TF = Transformer.from_crs("EPSG:4326", "EPSG:32612", always_xy=True)
cells_u = panel[["cell_id","lon","lat"]].drop_duplicates()
cell_xy = {}
for _, row in cells_u.iterrows():
    x, y = TF.transform(row["lon"], row["lat"])
    cell_xy[row["cell_id"]] = (x, y)
obs_xy = np.array([cell_xy[cid] for cid in panel["cell_id"]])


# ══════════════════════════════════════════════════════════════════════════════
# Conley (1999) spatial HAC
# ══════════════════════════════════════════════════════════════════════════════
def conley_se(result, obs_xy, cutoff_m=50_000):
    mu     = result.predict()
    y_arr  = np.asarray(result.model.endog)
    X_arr  = np.asarray(result.model.exog)
    scores = X_arr * (y_arr - mu)[:, None]
    XtWX   = X_arr.T @ (X_arr * mu[:, None])
    bread  = np.linalg.inv(XtWX)
    meat   = scores.T @ scores
    pairs  = cKDTree(obs_xy).query_pairs(cutoff_m, output_type="ndarray")
    if len(pairs):
        di_    = np.linalg.norm(obs_xy[pairs[:, 0]] - obs_xy[pairs[:, 1]], axis=1)
        k_vals = 1.0 - di_ / cutoff_m
        cross  = np.einsum("mi,mj,m->ij",
                           scores[pairs[:, 0]], scores[pairs[:, 1]], k_vals)
        meat  += cross + cross.T
    return np.sqrt(np.diag(bread @ meat @ bread))

def stars(p):
    return "***" if p < 0.01 else "**" if p < 0.05 else "*" if p < 0.1 else ""


# ══════════════════════════════════════════════════════════════════════════════
# Poisson DiD — 3 specifications
# ══════════════════════════════════════════════════════════════════════════════
print("\n" + "=" * 70)
print("Running Poisson DiD regressions...")
print("=" * 70)

y     = panel["deaths"]
vars1 = ["Post", "D_i", "Post_x_Di", "lat", "lon"]
vars2 = vars1 + ["dist_fence_km", "Post_x_fence", "dist_gap_km", "Post_x_gap"]

X1 = sm.add_constant(panel[vars1])
X2 = sm.add_constant(panel[vars2])

res1 = sm.Poisson(y, X1).fit(maxiter=500, disp=False)
con1 = conley_se(res1, obs_xy)
print("  Spec 1 done  (baseline)")

res2 = sm.Poisson(y, X2).fit(maxiter=500, disp=False)
con2 = conley_se(res2, obs_xy)
print("  Spec 2 done  (+ fence/gap distances)")

res3 = sm.Poisson(y, X2, offset=panel["log_apprehensions"]).fit(maxiter=500, disp=False)
con3 = conley_se(res3, obs_xy)
print("  Spec 3 done  (+ sector apprehension offset)")

LABELS = {
    "const"        : "Constant",
    "Post"         : "Post-SFA",
    "D_i"          : "Danger index (D_i)",
    "Post_x_Di"    : "Post × D_i  [funnel]",
    "lat"          : "Latitude",
    "lon"          : "Longitude",
    "dist_fence_km": "Dist to fence (km)",
    "Post_x_fence" : "Post × dist_fence",
    "dist_gap_km"  : "Dist to gap (km)",
    "Post_x_gap"   : "Post × dist_gap  [gap]",
}

all_vars = list(dict.fromkeys(
    list(res1.params.index) +
    [v for v in res2.params.index if v not in res1.params.index]
))

def safe(params, con_arr, pname):
    if pname not in params.index:
        return dict(beta=np.nan, IRR=np.nan, SE=np.nan, p=np.nan)
    i  = list(params.index).index(pname)
    b  = params[pname]; se = con_arr[i]
    z  = b / se; p = float(2 * (1 - stats.norm.cdf(abs(z))))
    return dict(beta=b, IRR=float(np.exp(b)), SE=se, p=p)

result_rows = []
for pname in all_vars:
    label = LABELS.get(pname, pname)
    r1 = safe(res1.params, con1, pname)
    r2 = safe(res2.params, con2, pname)
    r3 = safe(res3.params, con3, pname)
    result_rows.append(dict(variable=label,
        reg1_beta=r1["beta"], reg1_IRR=r1["IRR"], reg1_SE=r1["SE"], reg1_p=r1["p"],
        reg2_beta=r2["beta"], reg2_IRR=r2["IRR"], reg2_SE=r2["SE"], reg2_p=r2["p"],
        reg3_beta=r3["beta"], reg3_IRR=r3["IRR"], reg3_SE=r3["SE"], reg3_p=r3["p"],
    ))
reg = pd.DataFrame(result_rows)

print(f"\n  Key result — Post x D_i:")
for spec, r, c in [("1", res1, con1), ("2", res2, con2), ("3", res3, con3)]:
    d = safe(r.params, c, "Post_x_Di")
    if not np.isnan(d["beta"]):
        print(f"    Spec {spec}: β={d['beta']:.4f}  SE={d['SE']:.4f}  "
              f"p={d['p']:.4f} {stars(d['p'])}")


# ══════════════════════════════════════════════════════════════════════════════
# Word tables
# ══════════════════════════════════════════════════════════════════════════════
print("\n" + "=" * 70)
print("Building Word tables...")
print("=" * 70)

doc = Document()
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(11)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def center(cell):
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def bold_cell(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


# ── Table 1: DiD summary ──────────────────────────────────────────────────────
doc.add_heading(
    "Table 1. Migrant Deaths per Cell by Danger Index Tercile: Pre- vs. Post-SFA",
    level=2)

note1 = doc.add_paragraph(
    "Unit of observation is the grid cell-period (0.044 degree cells, ~2.7 miles per side). "
    "Danger index terciles are defined over the full sample of active cells using the "
    "composite Z-score (sum of 6 standardized factors: temperature, distance to city, "
    "distance to road, distance to water, slope, vegetation density). "
    "Difference = Post-SFA mean - Pre-SFA mean. "
    "% Change = Difference / Pre-SFA mean. "
    "t-statistics from two-sample t-tests (unequal variances). "
    "DiD row: columns show High-Low cross-difference pre- and post-SFA; "
    "% Change column shows difference in % changes (High % - Low %). "
    "*** p<0.01  ** p<0.05  * p<0.10."
)
note1.style.font.size = Pt(9)
note1.style.font.italic = True

# Assign terciles
cells_df = panel.drop_duplicates("cell_id")[["cell_id","D_i"]].copy()
cells_df["tercile"] = pd.qcut(cells_df["D_i"], q=3, labels=["Low","Medium","High"])
panel2 = panel.merge(cells_df[["cell_id","tercile"]], on="cell_id")
pre_p  = panel2[panel2["Post"] == 0]
post_p = panel2[panel2["Post"] == 1]

tercile_ranges = {}
for t in ["Low","Medium","High"]:
    v = cells_df[cells_df["tercile"] == t]["D_i"]
    tercile_ranges[t] = (v.min(), v.max())

t1 = doc.add_table(rows=1, cols=8)
t1.style = "Table Grid"
hdr = t1.rows[0].cells
for i, h in enumerate(["Danger Index\nTercile","Pre-SFA\nMean Deaths",
                        "Pre-SFA\nN Cells","Post-SFA\nMean Deaths",
                        "Post-SFA\nN Cells","Difference","% Change","t-statistic"]):
    hdr[i].text = h
    center(hdr[i]); bold_cell(hdr[i]); set_cell_bg(hdr[i], "D9D9D9")

for tier in ["Low","Medium","High"]:
    lo, hi     = tercile_ranges[tier]
    pre_vals   = pre_p[pre_p["tercile"]  == tier]["deaths"].values
    post_vals  = post_p[post_p["tercile"] == tier]["deaths"].values
    t_stat, pv = stats.ttest_ind(post_vals, pre_vals, equal_var=False)
    diff       = post_vals.mean() - pre_vals.mean()
    pct        = diff / pre_vals.mean() * 100
    row = t1.add_row().cells
    row[0].text = f"{tier} (D_i = {lo:.1f} to {hi:.1f})"
    row[1].text = f"{pre_vals.mean():.3f}"
    row[2].text = str(len(pre_vals))
    row[3].text = f"{post_vals.mean():.3f}"
    row[4].text = str(len(post_vals))
    row[5].text = f"{diff:+.3f}"
    row[6].text = f"{pct:+.1f}%"
    row[7].text = f"{t_stat:.2f}{stars(pv)}"
    for cell in row: center(cell)
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# DiD row
low_pre   = pre_p[pre_p["tercile"]  == "Low"]["deaths"].mean()
low_post  = post_p[post_p["tercile"] == "Low"]["deaths"].mean()
high_pre  = pre_p[pre_p["tercile"]  == "High"]["deaths"].mean()
high_post = post_p[post_p["tercile"] == "High"]["deaths"].mean()
pct_low   = (low_post  - low_pre)  / low_pre  * 100
pct_high  = (high_post - high_pre) / high_pre * 100
did_pct   = pct_high - pct_low
pre_diff  = high_pre  - low_pre
post_diff = high_post - low_post

panel_wide = panel2.pivot(index="cell_id", columns="Post", values="deaths").reset_index()
panel_wide.columns = ["cell_id","pre_d","post_d"]
panel_wide["change"] = panel_wide["post_d"] - panel_wide["pre_d"]
panel_wide = panel_wide.merge(cells_df[["cell_id","tercile"]], on="cell_id")
high_ch = panel_wide[panel_wide["tercile"] == "High"]["change"].values
low_ch  = panel_wide[panel_wide["tercile"] == "Low"]["change"].values
did_val = high_ch.mean() - low_ch.mean()
t_did, p_did = stats.ttest_ind(high_ch, low_ch, equal_var=False)

row = t1.add_row().cells
row[0].text = "DiD (High - Low)"
row[1].text = f"{pre_diff:+.3f}";  row[2].text = "H-L pre"
row[3].text = f"{post_diff:+.3f}"; row[4].text = "H-L post"
row[5].text = f"{did_val:+.3f}"
row[6].text = f"{did_pct:+.1f} pp"
row[7].text = f"{t_did:.2f}{stars(p_did)}"
for cell in row: center(cell); bold_cell(cell)
row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
set_cell_bg(row[0], "E6F0FF")

# All cells row
pre_all  = pre_p["deaths"].values
post_all = post_p["deaths"].values
t_all, p_all = stats.ttest_ind(post_all, pre_all, equal_var=False)
diff_all = post_all.mean() - pre_all.mean()
pct_all  = diff_all / pre_all.mean() * 100
row = t1.add_row().cells
row[0].text = "All cells"
row[1].text = f"{pre_all.mean():.3f}"; row[2].text = str(len(pre_all))
row[3].text = f"{post_all.mean():.3f}"; row[4].text = str(len(post_all))
row[5].text = f"{diff_all:+.3f}"
row[6].text = f"{pct_all:+.1f}%"
row[7].text = f"{t_all:.2f}{stars(p_all)}"
for cell in row: center(cell); bold_cell(cell)
row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
set_cell_bg(row[0], "F2F2F2")

interp1 = doc.add_paragraph(
    "Table 1 presents a simple difference-in-differences comparison of mean migrant deaths per "
    "grid cell across danger index terciles before and after the Secure Fence Act (SFA). "
    "Deaths increased in all three terciles following the SFA, but the magnitude differs sharply "
    "by terrain type. In the high-danger tercile, deaths per cell rose sharply and significantly, "
    "while the medium-danger tercile shows virtually no change. "
    "The difference-in-differences estimate — comparing the change in high-danger cells "
    "to the change in low-danger cells — is highly significant, confirming that the post-SFA "
    "increase was disproportionately concentrated in the most hazardous terrain. "
    "Notably, in the pre-SFA period the High-Low cross-difference is negative: high-danger "
    "cells had fewer deaths than low-danger cells before the SFA, because crossers were "
    "avoiding the most dangerous terrain. The SFA reversed this relationship. "
    "This pattern is consistent with a funnel effect: enforcement infrastructure channeled "
    "crossings away from urban ports of entry and into remote, high-danger corridors."
)
interp1.paragraph_format.space_before = Pt(6)
interp1.paragraph_format.space_after  = Pt(12)

doc.add_paragraph()


# ── Table 2: Regression results ────────────────────────────────────────────────
doc.add_heading(
    "Table 2. Poisson Difference-in-Differences: Migrant Deaths per Grid Cell",
    level=2)

note2 = doc.add_paragraph(
    "Poisson regression. Dependent variable: deaths per 0.044 degree grid cell-period. "
    "Specification (1): baseline DiD with danger index interaction. "
    "Specification (2): adds distance to nearest fence segment and fence gap. "
    "Specification (3): Specification (2) with log sector apprehensions as Poisson offset "
    "(dependent variable becomes death rate per crossing attempt). "
    "Standard errors are Conley (1999) spatial HAC with 50 km cutoff. "
    "*** p<0.01  ** p<0.05  * p<0.10."
)
note2.style.font.size = Pt(9)
note2.style.font.italic = True

VAR_LABELS = {
    "Constant"               : "Constant",
    "Post-SFA"               : "Post-SFA",
    "Danger index (D_i)"     : "Danger index (D_i)",
    "Post × D_i  [funnel]"   : "Post-SFA x D_i",
    "Latitude"               : "Latitude",
    "Longitude"              : "Longitude",
    "Dist to fence (km)"     : "Distance to fence (km)",
    "Post × dist_fence"      : "Post-SFA x Distance to fence",
    "Dist to gap (km)"       : "Distance to gap (km)",
    "Post × dist_gap  [gap]" : "Post-SFA x Distance to gap",
}

def fmt(beta, se, p):
    if pd.isna(beta):
        return "—", ""
    return f"{beta:.3f}{stars(p)}", f"({se:.3f})"

t2 = doc.add_table(rows=1, cols=7)
t2.style = "Table Grid"
hdr2 = t2.rows[0].cells
for i, h in enumerate(["Variable","(1) β","(1) SE","(2) β","(2) SE","(3) β","(3) SE"]):
    hdr2[i].text = h
    center(hdr2[i]); bold_cell(hdr2[i]); set_cell_bg(hdr2[i], "D9D9D9")

sub = t2.add_row().cells
sub[0].text = ""
sub[1].merge(sub[2]).text = "(1) Baseline"
sub[3].merge(sub[4]).text = "(2) + Fence/Gap"
sub[5].merge(sub[6]).text = "(3) + Offset"
for cell in [sub[0], sub[1], sub[3], sub[5]]:
    center(cell); bold_cell(cell); set_cell_bg(cell, "F2F2F2")

for _, r in reg.iterrows():
    label = VAR_LABELS.get(r["variable"], r["variable"])
    b1, s1 = fmt(r["reg1_beta"], r["reg1_SE"], r["reg1_p"])
    b2, s2 = fmt(r["reg2_beta"], r["reg2_SE"], r["reg2_p"])
    b3, s3 = fmt(r["reg3_beta"], r["reg3_SE"], r["reg3_p"])
    row = t2.add_row().cells
    row[0].text = label
    row[1].text = b1; row[2].text = s1
    row[3].text = b2; row[4].text = s2
    row[5].text = b3; row[6].text = s3
    for i in range(1, 7): center(row[i])
    if "Post-SFA x" in label:
        for cell in row: bold_cell(cell)

for label, v1, v2, v3 in [
    ("Observations",         f"{len(panel):,}", f"{len(panel):,}", f"{len(panel):,}"),
    ("Log-likelihood",       f"{res1.llf:.0f}", f"{res2.llf:.0f}", f"{res3.llf:.0f}"),
    ("AIC",                  f"{res1.aic:.0f}", f"{res2.aic:.0f}", f"{res3.aic:.0f}"),
    ("Pseudo-R2 (McFadden)", f"{res1.prsquared:.3f}", f"{res2.prsquared:.3f}", f"{res3.prsquared:.3f}"),
    ("Sector offset",        "No",  "No",  "Yes"),
    ("Conley SEs (50 km)",   "Yes", "Yes", "Yes"),
]:
    row = t2.add_row().cells
    row[0].text = label
    row[1].merge(row[2]).text = v1
    row[3].merge(row[4]).text = v2
    row[5].merge(row[6]).text = v3
    for cell in row:
        center(cell)
        for para in cell.paragraphs:
            for run in para.runs: run.italic = True
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

interp2 = doc.add_paragraph(
    "Table 2 presents Poisson regression estimates of the funnel effect across three "
    "specifications. The key coefficient of interest is Post-SFA x D_i, which captures whether "
    "cells with higher danger index values experienced a larger increase in deaths after the SFA. "
    "This coefficient is positive and statistically significant across all three specifications, "
    "indicating a robust funnel effect. The robustness of this result when the offset absorbs "
    "variation in crossing intensity suggests the effect reflects a genuine increase in the "
    "lethality of crossings in high-danger terrain, not merely an increase in crossing volume. "
    "Post-SFA x Distance to fence is negative and significant in specifications (2) and (3), "
    "consistent with the fence physically redirecting crossers into adjacent dangerous corridors. "
    "Post-SFA x Distance to gap is positive but falls short of conventional significance levels, "
    "suggesting proximity to fence gaps plays a secondary role once terrain danger is accounted for."
)
interp2.paragraph_format.space_before = Pt(6)
interp2.paragraph_format.space_after  = Pt(12)

doc.save(OUT_DOCX)
print(f"\nSaved: {OUT_DOCX}")
print("=" * 70)
print("Done.")
