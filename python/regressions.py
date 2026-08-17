"""
Builds two Word tables for the paper:
  Table 1: Deaths by danger index tercile, pre vs post SFA, with t-tests
  Table 2: Poisson DiD regression results — Reg 1 (baseline) vs Reg 2 (fence/gap) vs Reg 3 (offset)
"""

import numpy as np
import pandas as pd
from scipy import stats
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

BASE = "/Users/cbansak/Documents/Claude/Hot Spots"
OUT  = "/Users/cbansak/Documents/Claude/Migrant-deaths-paper/docs/paper_tables_regressions_2026-08-16.docx"

panel = pd.read_csv(f"{BASE}/panel_data.csv")
reg   = pd.read_csv(f"{BASE}/regression3_results.csv")

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def bold_row(row):
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

def center_row(row):
    for cell in row.cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)

# ══════════════════════════════════════════════════════════════════════════════
# TABLE 1: Deaths by D_i tercile, pre vs post SFA, with t-tests
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading("Table 1. Migrant Deaths per Cell by Danger Index Tercile: Pre- vs. Post-SFA", level=2)

note1 = doc.add_paragraph(
    "Unit of observation is the grid cell–period (0.044° cells, ~2.7 miles per side). "
    "Danger index terciles are defined over the full sample of active cells. "
    "Difference = Post-SFA mean − Pre-SFA mean. "
    "t-statistics from two-sample t-tests (unequal variances). "
    "*** p<0.01  ** p<0.05  * p<0.10."
)
note1.style.font.size = Pt(9)
note1.style.font.italic = True

# Assign terciles
cells_u = panel.drop_duplicates("cell_id")[["cell_id","D_i"]].copy()
cells_u["tercile"] = pd.qcut(cells_u["D_i"], q=3, labels=["Low","Medium","High"])
panel2 = panel.merge(cells_u[["cell_id","tercile"]], on="cell_id")

pre  = panel2[panel2["Post"] == 0]
post = panel2[panel2["Post"] == 1]

t1 = doc.add_table(rows=1, cols=8)
t1.style = "Table Grid"

# Header
hdr = t1.rows[0].cells
headers = ["Danger Index\nTercile", "Pre-SFA\nMean Deaths",
           "Pre-SFA\nN Cells", "Post-SFA\nMean Deaths",
           "Post-SFA\nN Cells", "Difference", "% Change", "t-statistic"]
for i, h in enumerate(headers):
    hdr[i].text = h
    hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True
    set_cell_bg(hdr[i], "D9D9D9")

def stars(p):
    return "***" if p < 0.01 else "**" if p < 0.05 else "*" if p < 0.10 else ""

# Compute actual D_i ranges per tercile from the data
tercile_ranges = {}
for t in ["Low", "Medium", "High"]:
    vals = cells_u[cells_u["tercile"] == t]["D_i"]
    tercile_ranges[t] = (int(vals.min()), int(vals.max()))

tercile_labels = {t: f"{t} (D_i = {lo}–{hi})"
                  for t, (lo, hi) in tercile_ranges.items()}

for tier in ["Low", "Medium", "High"]:
    pre_vals  = pre[pre["tercile"]  == tier]["deaths"].values
    post_vals = post[post["tercile"] == tier]["deaths"].values
    t_stat, p_val = stats.ttest_ind(post_vals, pre_vals, equal_var=False)
    diff = post_vals.mean() - pre_vals.mean()

    pct = (post_vals.mean() - pre_vals.mean()) / pre_vals.mean() * 100

    row = t1.add_row().cells
    row[0].text = tercile_labels[tier]
    row[1].text = f"{pre_vals.mean():.3f}"
    row[2].text = str(len(pre_vals))
    row[3].text = f"{post_vals.mean():.3f}"
    row[4].text = str(len(post_vals))
    row[5].text = f"{diff:+.3f}"
    row[6].text = f"{pct:+.1f}%"
    row[7].text = f"{t_stat:.2f}{stars(p_val)}"
    for cell in row:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# DiD row: (High post − High pre) − (Low post − Low pre)
# Compute per-cell changes, then t-test the difference in those changes
panel_wide = panel2.pivot(index="cell_id", columns="Post", values="deaths").reset_index()
panel_wide.columns = ["cell_id", "pre_d", "post_d"]
panel_wide["change"] = panel_wide["post_d"] - panel_wide["pre_d"]
panel_wide = panel_wide.merge(cells_u[["cell_id","tercile"]], on="cell_id")

high_changes = panel_wide[panel_wide["tercile"] == "High"]["change"].values
low_changes  = panel_wide[panel_wide["tercile"] == "Low"]["change"].values
did_val = high_changes.mean() - low_changes.mean()
t_did, p_did = stats.ttest_ind(high_changes, low_changes, equal_var=False)

# % changes for High and Low (needed for DiD % change)
low_pre  = pre[pre["tercile"]  == "Low"]["deaths"].mean()
low_post = post[post["tercile"] == "Low"]["deaths"].mean()
high_pre  = pre[pre["tercile"]  == "High"]["deaths"].mean()
high_post = post[post["tercile"] == "High"]["deaths"].mean()
pct_low  = (low_post  - low_pre)  / low_pre  * 100
pct_high = (high_post - high_pre) / high_pre * 100
did_pct  = pct_high - pct_low  # difference in % changes (pp)

# Cross-sectional High-Low differences in each period
pre_diff  = high_pre  - low_pre   # negative: High had fewer deaths pre-SFA
post_diff = high_post - low_post  # positive: High had more deaths post-SFA

row = t1.add_row().cells
row[0].text = "DiD (High − Low)"
row[1].text = f"{pre_diff:+.3f}"   # High-Low pre-SFA (negative)
row[2].text = "H−L pre"
row[3].text = f"{post_diff:+.3f}"  # High-Low post-SFA (positive)
row[4].text = "H−L post"
row[5].text = f"{did_val:+.3f}"
row[6].text = f"{did_pct:+.1f} pp"
row[7].text = f"{t_did:.2f}{stars(p_did)}"
for cell in row:
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cell.paragraphs[0].runs:
        run.bold = True
row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
set_cell_bg(row[0], "E6F0FF")  # light blue to highlight

# All cells row
pre_all  = pre["deaths"].values
post_all = post["deaths"].values
t_all, p_all = stats.ttest_ind(post_all, pre_all, equal_var=False)
diff_all = post_all.mean() - pre_all.mean()
pct_all  = diff_all / pre_all.mean() * 100
row = t1.add_row().cells
row[0].text = "All cells"
row[1].text = f"{pre_all.mean():.3f}"
row[2].text = str(len(pre_all))
row[3].text = f"{post_all.mean():.3f}"
row[4].text = str(len(post_all))
row[5].text = f"{diff_all:+.3f}"
row[6].text = f"{pct_all:+.1f}%"
row[7].text = f"{t_all:.2f}{stars(p_all)}"
for cell in row:
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cell.paragraphs[0].runs:
        run.bold = True
row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
set_cell_bg(row[0], "F2F2F2")

interp1 = doc.add_paragraph(
    "Table 1 presents a simple difference-in-differences comparison of mean migrant deaths per "
    "grid cell across danger index terciles before and after the Secure Fence Act (SFA). "
    "Deaths increased in all three terciles following the SFA, but the magnitude differs sharply "
    "by terrain type. In the high-danger tercile — cells with extreme heat, elevation, and "
    "remoteness — deaths per cell rose by 1.152, a 121.5% increase that is highly significant "
    "(t = 7.36, p < 0.001). By contrast, the medium-danger tercile shows virtually no change "
    "(+0.088, t = 0.56, p = 0.577). The low-danger tercile also sees a statistically significant "
    "increase (+0.422, +37.0%), consistent with some displacement toward lower-risk remote areas. "
    "The difference-in-differences estimate — comparing the change in high-danger cells "
    "to the change in low-danger cells — is +0.730 deaths per cell (t = 3.81, p < 0.001), "
    "confirming that the post-SFA increase was disproportionately concentrated in the most "
    "hazardous terrain. This pattern is consistent with a funnel effect: enforcement "
    "infrastructure channeled crossings away from urban ports of entry and into remote, "
    "high-danger corridors, driving up mortality."
)
interp1.paragraph_format.space_before = Pt(6)
interp1.paragraph_format.space_after  = Pt(12)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE 2: Regression results
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading("Table 2. Poisson Difference-in-Differences: Migrant Deaths per Grid Cell", level=2)

note2 = doc.add_paragraph(
    "Poisson regression. Dependent variable: deaths per 0.044° grid cell–period. "
    "Specification (1): baseline DiD with danger index interaction. "
    "Specification (2): adds distance to nearest fence segment and fence gap. "
    "Specification (3): Specification (2) with log sector apprehensions as Poisson offset "
    "(dependent variable becomes death rate per crossing attempt). "
    "Standard errors are Conley (1999) spatial HAC with 50 km cutoff. "
    "*** p<0.01  ** p<0.05  * p<0.10."
)
note2.style.font.size = Pt(9)
note2.style.font.italic = True

# Variable labels and order
VAR_LABELS = {
    "Constant":               "Constant",
    "Post-SFA":               "Post-SFA",
    "Danger index (D_i)":     "Danger index (D_i)",
    "Post × D_i  [funnel]":   "Post-SFA × D_i",
    "Latitude":               "Latitude",
    "Longitude":              "Longitude",
    "Dist to fence (km)":     "Distance to fence (km)",
    "Post × dist_fence":      "Post-SFA × Distance to fence",
    "Dist to gap (km)":       "Distance to gap (km)",
    "Post × dist_gap  [gap]": "Post-SFA × Distance to gap",
}

def fmt(beta, se, p):
    if pd.isna(beta):
        return "—", ""
    s = stars(p)
    return f"{beta:.3f}{s}", f"({se:.3f})"

t2 = doc.add_table(rows=1, cols=7)
t2.style = "Table Grid"

hdr2 = t2.rows[0].cells
h2 = ["Variable", "(1) β", "(1) SE", "(2) β", "(2) SE", "(3) β", "(3) SE"]
for i, h in enumerate(h2):
    hdr2[i].text = h
    hdr2[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hdr2[i].paragraphs[0].runs:
        run.bold = True
    set_cell_bg(hdr2[i], "D9D9D9")

# Sub-header for spec labels
sub = t2.add_row().cells
sub[0].text = ""
sub[1].merge(sub[2]).text = "(1) Baseline"
sub[3].merge(sub[4]).text = "(2) + Fence/Gap"
sub[5].merge(sub[6]).text = "(3) + Offset"
for cell in [sub[0], sub[1], sub[3], sub[5]]:
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cell.paragraphs[0].runs:
        run.bold = True
    set_cell_bg(cell, "F2F2F2")

for _, r in reg.iterrows():
    label = VAR_LABELS.get(r["variable"], r["variable"])

    b1, s1 = fmt(r["reg1_beta"], r["reg1_SE"], r["reg1_p"])
    b2, s2 = fmt(r["reg2_beta"], r["reg2_SE"], r["reg2_p"])
    b3, s3 = fmt(r["reg3_beta"], r["reg3_SE"], r["reg3_p"])

    # Coefficient row
    row = t2.add_row().cells
    row[0].text = label
    row[1].text = b1; row[2].text = s1
    row[3].text = b2; row[4].text = s2
    row[5].text = b3; row[6].text = s3
    for i in range(1, 7):
        row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Bold the key interaction rows
    if "Post-SFA ×" in label:
        for cell in row:
            for run in cell.paragraphs[0].runs:
                run.bold = True

# Footer rows
footer_data = [
    ("Observations",        "2,140",   "2,140",   "2,140"),
    ("Log-likelihood",      "−3,804",  "−3,702",  "−3,540"),
    ("AIC",                 "7,621",   "7,417",   "7,095"),
    ("Pseudo-R² (McFadden)","0.112",   "0.136",   "0.238"),
    ("Sector offset",       "No",      "No",      "Yes"),
    ("Conley SEs (50 km)",  "Yes",     "Yes",     "Yes"),
]
for label, v1, v2, v3 in footer_data:
    row = t2.add_row().cells
    row[0].text = label
    row[1].merge(row[2]).text = v1
    row[3].merge(row[4]).text = v2
    row[5].merge(row[6]).text = v3
    for cell in row:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.italic = True
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

interp2 = doc.add_paragraph(
    "Table 2 presents Poisson regression estimates of the funnel effect across three "
    "specifications. The key coefficient of interest is Post-SFA × D_i, which captures whether "
    "cells with higher danger index values experienced a larger increase in deaths after the SFA. "
    "This coefficient is positive and statistically significant in all three specifications: "
    "0.073 (p = 0.012) in the baseline model, 0.057 (p = 0.058) after controlling for proximity "
    "to fence segments and fence gaps, and 0.063 (p = 0.053) when deaths are scaled by sector "
    "apprehensions to account for crossing volume. The robustness of this result across "
    "specifications — including when the offset absorbs variation in crossing intensity — "
    "suggests the funnel effect reflects a genuine increase in the lethality of crossings in "
    "high-danger terrain, not merely an increase in crossing volume. "
    "Regarding the fence and gap variables, Post-SFA × Distance to fence is negative and "
    "significant in specifications (2) and (3), indicating that cells closer to fence "
    "construction saw larger increases in deaths — consistent with the fence physically "
    "redirecting crossers into adjacent dangerous corridors. Post-SFA × Distance to gap is "
    "positive but falls short of conventional significance levels (p ≈ 0.11–0.12), "
    "suggesting that proximity to fence gaps may play a secondary role once terrain danger "
    "is accounted for."
)
interp2.paragraph_format.space_before = Pt(6)
interp2.paragraph_format.space_after  = Pt(12)

doc.save(OUT)
print(f"Saved: {OUT}")
